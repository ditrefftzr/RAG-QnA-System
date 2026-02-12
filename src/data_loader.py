"""
Document loader for RAG system
Loads CV, projects, certificates, and other documents with metadata
"""
import os
import sys
from pathlib import Path
from typing import List, Dict
from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
    DirectoryLoader
)
from langchain_core.documents import Document

# Add parent directory to path to import config
sys.path.insert(0, str(Path(__file__).parent.parent))
from config import get_logger, PATHS

# Create logger for this module
logger = get_logger(__name__)


class DocumentLoader:
    """Loads documents from the data/raw directory with metadata"""
    
    def __init__(self, data_dir: str = None):
        """
        Initialize the document loader
        
        Args:
            data_dir: Path to directory containing raw documents (uses config default if None)
        """
        # Use config default if not provided
        if data_dir is None:
            data_dir = PATHS['data_directory']
        
        self.data_dir = Path(data_dir)
        
        if not self.data_dir.exists():
            raise ValueError(f"Data directory not found: {data_dir}")
    
    def load_all_documents(self) -> List[Document]:
        """
        Load all documents from data directory
        
        Returns:
            List of Document objects with content and metadata
        """
        documents = []
        
        logger.info(f"Loading documents from: {self.data_dir}")
        
        # Load PDFs
        pdf_docs = self._load_pdfs()
        documents.extend(pdf_docs)
        logger.info(f"Loaded {len(pdf_docs)} PDF documents")
        
        # Load text files
        txt_docs = self._load_text_files()
        documents.extend(txt_docs)
        logger.info(f"Loaded {len(txt_docs)} text documents")
        
        # Load SQL files
        sql_docs = self._load_sql_files()
        documents.extend(sql_docs)
        logger.info(f"Loaded {len(sql_docs)} SQL documents")

        # Load docx files
        docx_docs = self._load_docx_files()
        documents.extend(docx_docs)
        logger.info(f"Loaded {len(docx_docs)} Word documents")
        
        logger.info(f"Total documents loaded: {len(documents)}")
        
        return documents
    
    def _load_pdfs(self) -> List[Document]:
        """Load all PDF files from data directory"""
        pdf_documents = []
        
        # Find all PDF files
        pdf_files = list(self.data_dir.glob("*.pdf"))
        
        for pdf_path in pdf_files:
            try:
                loader = PyPDFLoader(str(pdf_path))
                docs = loader.load()
                
                # Add custom metadata
                for doc in docs:
                    doc.metadata.update({
                        "source_file": pdf_path.name,
                        "file_type": "pdf",
                        "category": self._categorize_document(pdf_path.name)
                    })
                
                pdf_documents.extend(docs)
                logger.debug(f"{pdf_path.name}: {len(docs)} pages")
                
            except Exception as e:
                logger.warning(f"Error loading {pdf_path.name}: {e}")
        
        return pdf_documents
    
    def _load_text_files(self) -> List[Document]:
        """Load all text files from data directory"""
        text_documents = []
        
        # Find all text files (excluding .sql)
        txt_files = [f for f in self.data_dir.glob("*.txt")]
        
        for txt_path in txt_files:
            try:
                loader = TextLoader(str(txt_path), encoding='utf-8')
                docs = loader.load()
                
                # Add custom metadata
                for doc in docs:
                    doc.metadata.update({
                        "source_file": txt_path.name,
                        "file_type": "txt",
                        "category": self._categorize_document(txt_path.name)
                    })
                
                text_documents.extend(docs)
                logger.debug(f"{txt_path.name}")
                
            except Exception as e:
                logger.warning(f"Error loading {txt_path.name}: {e}")
        
        return text_documents
    
    def _load_sql_files(self) -> List[Document]:
        """Load all SQL files from data directory"""
        sql_documents = []
        
        # Find all SQL files
        sql_files = list(self.data_dir.glob("*.sql"))
        
        for sql_path in sql_files:
            try:
                # Read SQL file with UTF-8 encoding
                with open(sql_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                # Create Document object
                doc = Document(
                    page_content=content,
                    metadata={
                        "source_file": sql_path.name,
                        "file_type": "sql",
                        "category": "technical_skills"  # SQL files show database skills
                    }
                )
                
                sql_documents.append(doc)
                logger.debug(f"{sql_path.name}")
                
            except Exception as e:
                logger.warning(f"Error loading {sql_path.name}: {e}")
        
        return sql_documents
    
    def _load_docx_files(self) -> List[Document]:
        """
        Load Word documents (.docx) from data directory
        
        Returns:
            List of Document objects from Word files
        """
        from docx import Document as DocxDocument
        
        documents = []
        docx_files = list(self.data_dir.glob("*.docx"))
        
        for docx_path in docx_files:
            try:
                # Load Word document
                doc = DocxDocument(docx_path)
                
                # Extract all text from paragraphs
                full_text = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():  # Skip empty paragraphs
                        full_text.append(paragraph.text)
                
                # Combine into single text
                content = "\n".join(full_text)
                
                # Create Document object
                category = self._categorize_document(docx_path.name)
                
                doc_obj = Document(
                    page_content=content,
                    metadata={
                        "source_file": docx_path.name,
                        "category": category,
                        "file_type": "docx"
                    }
                )
                
                documents.append(doc_obj)
                logger.debug(f"{docx_path.name}: {len(doc.paragraphs)} paragraphs")
                
            except Exception as e:
                logger.warning(f"Error loading {docx_path.name}: {e}")
        
        return documents

    
    def _categorize_document(self, filename: str) -> str:
        """
        Categorize document based on filename
        
        Args:
            filename: Name of the file
            
        Returns:
            Category string (cv, project, certificate, academic, experience, technical_skills)
        """
        filename_lower = filename.lower()
        
        if "cv" in filename_lower or "resume" in filename_lower:
            return "cv"
        elif "certificate" in filename_lower or "certification" in filename_lower:
            return "certificate"
        elif any(word in filename_lower for word in ["mysql", "powerbi", "power_bi", "power bi"]):
            return "certificate"
        elif any(word in filename_lower for word in ["proyecto", "project", "integrador"]):
            return "project"
        elif any(word in filename_lower for word in ["saber", "resultados", "academic"]):
            return "academic"
        elif any(word in filename_lower for word in ["pln", "nlp", "mlops", "clasificacion", "classification"]):
            return "project"
        elif "database" in filename_lower or "sql" in filename_lower:
            return "technical_skills"
        else:
            return "other"
    
    def get_document_stats(self, documents: List[Document]) -> Dict:
        """
        Get statistics about loaded documents
        
        Args:
            documents: List of loaded documents
            
        Returns:
            Dictionary with statistics
        """
        stats = {
            "total_documents": len(documents),
            "by_category": {},
            "by_file_type": {},
            "total_characters": 0
        }
        
        for doc in documents:
            # Count by category
            category = doc.metadata.get("category", "unknown")
            stats["by_category"][category] = stats["by_category"].get(category, 0) + 1
            
            # Count by file type
            file_type = doc.metadata.get("file_type", "unknown")
            stats["by_file_type"][file_type] = stats["by_file_type"].get(file_type, 0) + 1
            
            # Total characters
            stats["total_characters"] += len(doc.page_content)
        
        return stats


# Test function
if __name__ == "__main__":
    """Test the document loader"""
    print("=" * 80)
    print("TESTING DOCUMENT LOADER")
    print("=" * 80)
    
    # Initialize loader
    loader = DocumentLoader()
    
    # Load all documents
    documents = loader.load_all_documents()
    
    # Get statistics
    stats = loader.get_document_stats(documents)
    
    print("\n" + "=" * 80)
    print("DOCUMENT STATISTICS")
    print("=" * 80)
    print(f"Total documents: {stats['total_documents']}")
    print(f"Total characters: {stats['total_characters']:,}")
    print(f"\nBy category:")
    for category, count in stats['by_category'].items():
        print(f"  • {category}: {count}")
    print(f"\nBy file type:")
    for file_type, count in stats['by_file_type'].items():
        print(f"  • {file_type}: {count}")
    
    # Show sample documents from each type
    print("\n" + "=" * 80)
    print("SAMPLE DOCUMENTS BY TYPE")
    print("=" * 80)
    
    file_types_shown = set()
    for doc in documents:
        file_type = doc.metadata.get('file_type')
        if file_type not in file_types_shown:
            print(f"\n{file_type.upper()} - {doc.metadata.get('source_file')}:")
            print(f"Category: {doc.metadata.get('category')}")
            print(f"Content preview (first 150 chars):")
            print(doc.page_content[:150] + "...")
            file_types_shown.add(file_type)