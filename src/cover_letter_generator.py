"""
Cover Letter Generator
Automatically generates customized cover letters from job descriptions
"""
import os
import sys
from pathlib import Path
from typing import Dict, List, Optional
from dotenv import load_dotenv
from google import genai
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import json

# Add parent directory to path to import config
sys.path.insert(0, str(Path(__file__).parent.parent))
from config import get_logger, MODEL_CONFIG, PATHS, GENERATION_CONFIG

from src.rag_chain import RAGChain
from src.vector_store import VectorStoreManager

load_dotenv()

# Create logger for this module
logger = get_logger(__name__)


class CoverLetterGenerator:
    """Generates customized cover letters based on job descriptions"""
    
    def __init__(self, rag_chain: RAGChain):
        """
        Initialize the cover letter generator
        
        Args:
            rag_chain: RAG chain for retrieving relevant experience
        """
        self.rag = rag_chain
        
        # Initialize Gemini for job description analysis
        api_key = os.getenv("GOOGLE_API_KEY")
        self.client = genai.Client(api_key=api_key)
        self.model_name = MODEL_CONFIG['generation_model']
        
        logger.info(f"Cover Letter Generator initialized with model: {self.model_name}")
    
    def extract_job_info(self, job_description: str) -> Dict[str, any]:
        """
        Extract structured information from job description
        
        Args:
            job_description: Full job posting text
            
        Returns:
            Dictionary with company, role, requirements, responsibilities
        """
        logger.info("Analyzing job description...")
        
        extraction_prompt = f"""
You are an information extraction specialist. Your task is to parse 
job postings and extract structured data with high accuracy. Focus on 
extracting exactly what's written, without interpretation.

CRITICAL INSTRUCTIONS:
- Return ONLY valid JSON, no markdown, no code blocks, no explanations
- Extract exactly what's written, don't infer or add information
- If something is not mentioned, use null

Job Description:
{job_description}

Return a JSON object with this EXACT structure:
{{
    "company_name": "Company name or null",
    "job_title": "Job title/role or null",
    "must_have_skills": ["skill1", "skill2", ...],
    "nice_to_have_skills": ["skill1", "skill2", ...],
    "responsibilities": ["responsibility1", "responsibility2", ...],
    "key_requirements": ["requirement1", "requirement2", ...]
}}

JSON:"""

        try:
            response = self.client.models.generate_content(
                model=self.model_name,
                contents=extraction_prompt,
                config={
                    'temperature': GENERATION_CONFIG['temperature'],
                    'max_output_tokens': GENERATION_CONFIG['max_tokens'],
                    'top_p': GENERATION_CONFIG['top_p'],
                }
            )
            
            # Get the response text
            response_text = response.text.strip()
            
            # Remove markdown code blocks if present
            if response_text.startswith('```'):
                response_text = response_text.split('```')[1]
                if response_text.startswith('json'):
                    response_text = response_text[4:]
                response_text = response_text.strip()
            
            # Parse JSON
            job_info = json.loads(response_text)
            
            logger.info("Extracted job info:")
            logger.info(f"  Company: {job_info.get('company_name', 'Unknown')}")
            logger.info(f"  Role: {job_info.get('job_title', 'Unknown')}")
            logger.info(f"  Must-have skills: {len(job_info.get('must_have_skills', []))}")
            logger.info(f"  Nice-to-have skills: {len(job_info.get('nice_to_have_skills', []))}")
            
            return job_info
            
        except Exception as e:
            logger.error(f"Error extracting job info: {e}")
            # Return minimal structure if extraction fails
            return {
                "company_name": None,
                "job_title": None,
                "must_have_skills": [],
                "nice_to_have_skills": [],
                "responsibilities": [],
                "key_requirements": []
            }


    def match_requirements_to_experience(self, job_info: Dict) -> Dict[str, any]:
        """
        Match job requirements to candidate's actual experience
        Uses a single comprehensive query for efficiency
        
        Args:
            job_info: Extracted job information
            
        Returns:
            Dictionary with verified skills and supporting evidence
        """
        logger.info("Matching requirements to your experience...")
        
        must_have = job_info.get('must_have_skills', [])
        nice_to_have = job_info.get('nice_to_have_skills', [])
        all_skills = must_have + nice_to_have
        
        logger.info(f"Checking {len(all_skills)} skills...")
        
        # Build single comprehensive query
        skills_text = "\n".join([f"- {skill}" for skill in all_skills])
        
        query = f"""Review my professional background documents (CV, resume, work history, certifications, and project portfolio) to identify which of these job requirements I meet:

        {skills_text}

        Important: Search ALL document types including:
        - CV/Resume (education, work history, technical skills, certifications)
        - Employment records and job descriptions
        - Project documentation and technical writeups
        - Database schemas and code repositories
        - Training certificates and coursework

        For each requirement I have experience with, cite specific evidence.
        Do not include requirements without clear documentation."""

        # Single RAG call to get comprehensive answer
        logger.debug("Searching your background...")
        result = self.rag.generate_answer(
            question=query,
            purpose="general",
            tone="professional",
            k = max(5, min(len(all_skills), 12)) ## Custom amount of documents returned based on the length 
        )
        
        answer = result['answer']
        logger.info("Analysis complete")

        # DEBUG: Print the analysis
        logger.debug(f"Analysis from LLM: {answer[:500]}...")

        # Extract structured list of verified skills
        logger.debug("Extracting verified skills...")
        
        extraction_prompt = f"""You are extracting verified skills from an analysis.

        ORIGINAL SKILL LIST:
        {skills_text}

        ANALYSIS:
        {answer}

        TASK:
        From the ORIGINAL SKILL LIST above, identify which skills are mentioned as confirmed/verified in the ANALYSIS.

        Look for skills that have:
        - Specific evidence mentioned
        - Project names or work experience cited
        - Concrete examples provided

        IMPORTANT:
        - Match the EXACT skill names from the original list
        - Include partial matches (e.g., if "machine learning" is in the list and analysis discusses ML projects, include it)
        - Be thorough - include all skills with evidence

        Return ONLY a JSON array with the verified skill names:
        ["skill1", "skill2", ...]

JSON:"""

        try:
            response = self.client.models.generate_content(
                model=self.model_name,
                contents=extraction_prompt,
                config={
                    'temperature': GENERATION_CONFIG['temperature'],
                    'max_output_tokens': GENERATION_CONFIG['max_tokens'],
                    'top_p': GENERATION_CONFIG['top_p'],
                }
            )
            
            response_text = response.text.strip()
            
            # Clean markdown if present
            if response_text.startswith('```'):
                response_text = response_text.split('```')[1]
                if response_text.startswith('json'):
                    response_text = response_text[4:]
                response_text = response_text.strip()
            
            verified_skills = json.loads(response_text)
            
            logger.info(f"Verified {len(verified_skills)} out of {len(all_skills)} skills")
            
            return {
                'verified_skills': verified_skills,
                'supporting_evidence': answer,  # Full narrative with examples
                'sources': result['sources'],
                'must_have_matched': [s for s in verified_skills if s in must_have],
                'nice_to_have_matched': [s for s in verified_skills if s in nice_to_have]
            }
            
        except Exception as e:
            logger.error(f"Error parsing verified skills: {e}")
            logger.debug(f"Raw response: {response_text[:200]}...")
            
            # Fallback: return with unparsed evidence
            return {
                'verified_skills': [],
                'supporting_evidence': answer,
                'sources': result['sources'],
                'must_have_matched': [],
                'nice_to_have_matched': []
            }
        
    def generate_cover_letter(
        self,
        job_description: str,
        output_dir: str = None
    ) -> dict:
        """
        Generate a customized cover letter for a job application.
        Outputs as .DOCX with Times New Roman formatting.
        
        Args:
            job_description: Full text of the job posting
            output_dir: Directory to save the cover letter (uses config default if None)
            
        Returns:
            dict with:
                - cover_letter_text: The generated text
                - file_path: Path to saved .docx file
                - metadata: Job info and match statistics
        """
        # Use config default if not provided
        if output_dir is None:
            output_dir = PATHS['cover_letters_directory']
        
        logger.info("="*80)
        logger.info("COVER LETTER GENERATOR")
        logger.info("="*80)
        
        # Step 1: Extract job information
        logger.info("Step 1: Extracting job information...")
        job_info = self.extract_job_info(job_description)
        
        # Extract values with safe defaults
        company = job_info.get('company_name') or 'Not specified'
        role = job_info.get('job_title') or 'Not specified'
        must_have_skills = job_info.get('must_have_skills', [])
        nice_to_have_skills = job_info.get('nice_to_have_skills', [])
        
        if company == 'Not specified':
            logger.warning("Company name not detected in job posting")
        
        logger.info(f"Company: {company}")
        logger.info(f"Role: {role}")
        logger.info(f"Requirements extracted:")
        logger.info(f"  Must-have: {len(must_have_skills)}")
        logger.info(f"  Nice-to-have: {len(nice_to_have_skills)}")
        
        # Step 2: Match requirements to experience
        logger.info("Step 2: Matching requirements to your experience...")
        matched_experience = self.match_requirements_to_experience(job_info)
        
        # Extract match results
        verified_skills = matched_experience.get('verified_skills', [])
        must_have_matched = matched_experience.get('must_have_matched', [])
        nice_to_have_matched = matched_experience.get('nice_to_have_matched', [])
        supporting_evidence = matched_experience.get('supporting_evidence', '')
        
        logger.info(f"Must-have skills matched: {len(must_have_matched)}/{len(must_have_skills)}")
        logger.info(f"Nice-to-have skills matched: {len(nice_to_have_matched)}/{len(nice_to_have_skills)}")
        
        # Step 3: Generate the cover letter text
        logger.info("Step 3: Generating cover letter...")

        # Format verified skills for prompt
        skills_list = ", ".join(verified_skills[:5])  # Top 5 skills

        # Determine company reference strategy
        if company and company != "Not specified":
            company_instruction = f'Use the company name "{company}" directly throughout the letter.'
            company_reference = company
        else:
            company_instruction = 'Since the company name is not specified, use natural alternatives like "your team", "your organization", "this opportunity", or refer to the role/work directly.'
            company_reference = "your team"

        prompt = f"""You are writing a compelling cover letter for a job application.

        JOB DETAILS:
        - Company: {company_reference}
        - Position: {role}

        CANDIDATE'S VERIFIED SKILLS AND EXPERIENCE:
        {supporting_evidence}

        VERIFIED SKILLS TO HIGHLIGHT:
        {skills_list}

        INSTRUCTIONS:
        Write a 300-word cover letter with the following structure:

        1. OPENING (2-3 sentences):
        - Express genuine enthusiasm for the {role} position
        - Immediately state your unique value proposition
        - Mention your background that makes you a strong fit

        2. EXPERIENCE ALIGNMENT (Main body - 4-5 sentences):
        - Highlight VERIFIED skills with specific examples from the evidence
        - Use concrete metrics and achievements (project names, percentages, scale)
        - Show natural alignment with the role requirements
        - Emphasize unique combination: Finance + Data Science + GenAI

        3. WHY THIS ROLE (2-3 sentences):
        - Express what excites you about this opportunity and the type of work
        - Connect your interests and values to the work described in the job posting
        - Focus on the impact you want to make and problems you want to solve

        4. CLOSING (1-2 sentences):
        - Reiterate enthusiasm
        - Clear call to action
        - Professional but warm

        CRITICAL RULES - FOLLOW EXACTLY:

        TONE & STYLE:
        ✅ Write in first-person, enthusiastic, warm tone
        ✅ Use natural, conversational language - not templated or robotic
        ✅ Be specific and concrete with examples
        ✅ Target exactly 300 words

        CONTENT ACCURACY:
        ✅ ONLY mention skills that appear in the VERIFIED SKILLS list
        ✅ NEVER mention skills, technologies, or experience not in the verified list
        ✅ Use actual project names, metrics, and achievements from the evidence
        ✅ {company_instruction}

        CONFIDENCE & POSITIONING - CRITICAL:
        ✅ Present ALL verified skills with CONFIDENCE - you have real experience with them
        ✅ Use assertive language: "I have experience with", "I've built", "I've worked with", "I'm ready to", "I'm prepared to"
        ✅ Focus on what you CAN DO and HAVE DONE, not what you want to learn
        ✅ Treat ALL verified skills as ESTABLISHED EXPERTISE, not emerging/developing knowledge

        ❌ NEVER use uncertain/junior language - BANNED PHRASES:
        ❌ "I'm eager to learn"
        ❌ "I'm eager to apply"
        ❌ "I'm eager to expand"
        ❌ "I'm eager to discuss"
        ❌ "I'm eager to contribute"
        ❌ "I'm willing to expand my knowledge"
        ❌ "I look forward to developing my skills in"
        ❌ "I'm excited to grow my experience with"
        ❌ "I hope to apply and expand"
        ❌ "I'm interested in learning more about"
        ❌ "I would like to gain more experience in"
        ❌ "at your earliest convenience"

        ❌ NEVER use diminishing qualifiers on verified skills:
        ❌ "emerging knowledge in"
        ❌ "emerging experience in"
        ❌ "developing experience with"
        ❌ "growing expertise in"
        ❌ "budding skills in"
        ❌ "nascent experience with"
        ❌ "preliminary knowledge of"
        ❌ "foundational understanding of"
        ❌ "basic familiarity with"

        INSTEAD - Use confident, established language:
        ✅ "I have experience with"
        ✅ "and expertise in"
        ✅ "with skills in"
        ✅ "including experience with"
        ✅ "I've worked with"
        ✅ "I'm ready to apply"
        ✅ "I can apply"
        ✅ "I'm prepared to"
        ✅ "I'd welcome the opportunity to discuss"

        CRITICAL PRINCIPLE:
        If a skill is in the VERIFIED SKILLS list, it is ESTABLISHED EXPERTISE.
        Present it confidently and directly - no hedging, no qualifiers, no "emerging".
        You HAVE the skill. Period.

        FORBIDDEN ELEMENTS - NEVER INCLUDE:
        ❌ Square brackets [] anywhere: no [Company Name], [mention X], [add details]
        ❌ Meta-instructions: no "research this!", "elaborate here", "add specifics"
        ❌ Parenthetical notes: no "(mention achievement)", "(add details)"
        ❌ Placeholder text of any kind
        ❌ Direct quotes from the job description
        ❌ Signature block or "Sincerely" - just the letter body
        ❌ Incomplete sentences or gaps requiring fill-in

        SPECIFIC EXAMPLES - DO'S AND DON'TS:

        Company References:
        ❌ BAD: "I'm excited about [Company Name]'s mission"
        ❌ BAD: "contributing to [company initiative - research this]"
        ✅ GOOD (if company known): "I'm excited about Booth's commitment to responsible AI"
        ✅ GOOD (if company unknown): "I'm excited about this opportunity to build production AI systems"

        Technical Skills:
        ❌ BAD: "My experience with [specific technology]"
        ❌ BAD: "[mention relevant project here]"
        ✅ GOOD: "My experience developing RAG pipelines with LangChain and ChromaDB"
        ✅ GOOD: "I built a sentiment analysis system achieving 85% accuracy"

        Achievements:
        ❌ BAD: "I improved performance by [X]%"
        ❌ BAD: "(add metric here)"
        ✅ GOOD: "I improved model accuracy by 15-20 percentage points"
        ✅ GOOD: "I built infrastructure serving 200+ users with 80% efficiency gains"

        WRITING STRATEGY:
        - If you lack specific company information, focus on what excites you about the ROLE, the WORK, and the IMPACT
        - Write complete, confident statements - no hedging or gaps
        - Every sentence should add value - no filler
        - Be authentic and specific, not generic
        - Present yourself as READY TO CONTRIBUTE NOW, not as someone who needs training

        Write the cover letter now (body only, no signature):"""

        response = self.client.models.generate_content(
            model=self.model_name,
            contents=prompt,
            config={
                'temperature': GENERATION_CONFIG['temperature'],
                'max_output_tokens': GENERATION_CONFIG['max_tokens'],
                'top_p': GENERATION_CONFIG['top_p'],
            }
        )
        cover_letter_text = response.text.strip()

        word_count = len(cover_letter_text.split())
        logger.info(f"Generated! Word count: {word_count}")
        
        # Step 4: Create Word document with Times New Roman formatting
        logger.info("Step 4: Creating Word document with Times New Roman...")
        
        # Create a new Document
        doc = Document()
        
        # Set up the document margins (1 inch all around - standard for cover letters)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Helper function to add formatted text (reduces code repetition)
        def add_formatted_paragraph(text, space_after=0):
            """Add a paragraph with Times New Roman 12pt formatting"""
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            if space_after:
                para.paragraph_format.space_after = Pt(space_after)
            return para
        
        # Contact information
        add_formatted_paragraph("David Trefftz Restrepo")
        doc.add_paragraph()  # blank line
        
        # Date
        add_formatted_paragraph(datetime.now().strftime("%B %d, %Y"))
        doc.add_paragraph()  # blank line
        
        # Company info (if available)
        if company and company != "Not specified":
            add_formatted_paragraph(f"Hiring Manager\n{company}")
            doc.add_paragraph()
        
        # Cover letter body (split by paragraphs to maintain structure)
        paragraphs = cover_letter_text.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():  # Only add non-empty paragraphs
                add_formatted_paragraph(para_text.strip(), space_after=12)
        
        # Signature
        doc.add_paragraph()
        add_formatted_paragraph("Sincerely,\n\nDavid Trefftz Restrepo")
        
        # Step 5: Save the document
        logger.info("Step 5: Saving Word document...")
        
        # Create output directory if it doesn't exist
        os.makedirs(output_dir, exist_ok=True)
        
        # Generate filename
        company_slug = company.replace(' ', '_').replace('/', '_')
        role_slug = role.replace(' ', '_').replace('/', '_')
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{company_slug}_{role_slug}_{timestamp}.docx"
        file_path = os.path.join(output_dir, filename)
        
        # Save the Word document
        doc.save(file_path)
        logger.info(f"Saved to: {file_path}")
        
        # Step 6: Prepare metadata
        metadata = {
            'company': company,
            'role': role,
            'generated_date': timestamp,
            'word_count': word_count,
            'must_have_matches': len(must_have_matched),
            'must_have_total': len(must_have_skills),
            'nice_to_have_matches': len(nice_to_have_matched),
            'nice_to_have_total': len(nice_to_have_skills),
        }
        
        logger.info("="*80)
        logger.info("COVER LETTER GENERATION COMPLETE!")
        logger.info("="*80)
        logger.info("Match Statistics:")
        logger.info(f"  Must-have requirements: {metadata['must_have_matches']}/{metadata['must_have_total']}")
        logger.info(f"  Nice-to-have requirements: {metadata['nice_to_have_matches']}/{metadata['nice_to_have_total']}")
        logger.info(f"Word document saved to: {file_path}")
        
        return {
            'cover_letter_text': cover_letter_text,
            'file_path': file_path,
            'metadata': metadata
        }


if __name__ == "__main__":
    """Test the cover letter generator"""
    
    print("=" * 80)
    print("COVER LETTER GENERATOR - TEST")
    print("=" * 80)
    
    # Initialize
    print("\n[Initializing System]")
    vs_manager = VectorStoreManager()
    vs_manager.load_vectorstore()
    rag = RAGChain(vs_manager)
    
    generator = CoverLetterGenerator(rag)
    
    # Test job description
    sample_job_description = """
    Data Scientist - Bancolombia
    
    Bancolombia, Colombia's leading financial institution, is seeking a Data Scientist to join our AI Innovation team.
    
    Responsibilities:
    - Build and deploy machine learning models for credit risk assessment
    - Analyze large financial datasets to extract insights
    - Collaborate with product teams to implement AI solutions
    - Develop NLP models for document processing
    
    Requirements:
    - 2+ years experience with Python and machine learning
    - Strong SQL skills and database knowledge
    - Experience with financial data analysis
    - Proven track record of deploying ML models to production
    
    Nice to have:
    - NLP experience with transformers (BERT, RoBERTa)
    - Knowledge of MLOps practices
    - Experience with cloud platforms (AWS, GCP)
    """
    
    # Generate full cover letter
    result = generator.generate_cover_letter(sample_job_description)
    
    print("\n📄 GENERATED COVER LETTER:")
    print("=" * 80)
    print(result['cover_letter_text'])
    print("=" * 80)
    
    print(f"\n💾 Saved to: {result['file_path']}")
    
    print("\n" + "=" * 80)
    print("✅ ALL TESTS COMPLETE")
    print("=" * 80)